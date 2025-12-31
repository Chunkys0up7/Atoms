import json
import os
import re
import signal
import subprocess
from collections import defaultdict
from datetime import datetime, timedelta
from pathlib import Path
from typing import Any, Dict, List, Optional

import psutil
from fastapi import APIRouter, HTTPException
from fastapi.responses import JSONResponse

router = APIRouter()

# Global MkDocs process reference
_mkdocs_process: Optional[subprocess.Popen] = None
MKDOCS_PORT = 8001
MKDOCS_HOST = "127.0.0.1"


def get_mkdocs_dir() -> Path:
    """Get the docs directory containing mkdocs.yml"""
    # MkDocs config is in the project root
    base = Path(__file__).parent.parent.parent
    return base


def is_mkdocs_running() -> bool:
    """Check if MkDocs server is running on the configured port."""
    # Check if we have a process reference and it's still running
    if _mkdocs_process and _mkdocs_process.poll() is None:
        return True

    # Check if any process is using the MkDocs port
    for conn in psutil.net_connections(kind="inet"):
        if conn.laddr.port == MKDOCS_PORT and conn.status == "LISTEN":
            return True

    return False


def get_mkdocs_pid() -> Optional[int]:
    """Get the PID of the running MkDocs process."""
    if _mkdocs_process and _mkdocs_process.poll() is None:
        return _mkdocs_process.pid

    # Try to find the process by port
    for conn in psutil.net_connections(kind="inet"):
        if conn.laddr.port == MKDOCS_PORT and conn.status == "LISTEN":
            return conn.pid

    return None


@router.get("/api/mkdocs/status")
def get_mkdocs_status() -> Dict[str, Any]:
    """Get the current status of the MkDocs server."""
    running = is_mkdocs_running()
    pid = get_mkdocs_pid() if running else None

    return {
        "running": running,
        "pid": pid,
        "url": f"http://{MKDOCS_HOST}:{MKDOCS_PORT}" if running else None,
        "port": MKDOCS_PORT,
        "host": MKDOCS_HOST,
    }


@router.post("/api/mkdocs/start")
def start_mkdocs() -> Dict[str, Any]:
    """Start the MkDocs development server."""
    global _mkdocs_process

    if is_mkdocs_running():
        return {
            "status": "already_running",
            "message": f"MkDocs server already running on port {MKDOCS_PORT}",
            "url": f"http://{MKDOCS_HOST}:{MKDOCS_PORT}",
        }

    mkdocs_dir = get_mkdocs_dir()

    if not (mkdocs_dir / "mkdocs.yml").exists():
        raise HTTPException(status_code=500, detail=f"mkdocs.yml not found in {mkdocs_dir}")

    try:
        # Start MkDocs server
        _mkdocs_process = subprocess.Popen(
            [
                "mkdocs",
                "serve",
                "--dev-addr",
                f"{MKDOCS_HOST}:{MKDOCS_PORT}",
                "--no-livereload",  # Disable livereload to prevent conflicts
            ],
            cwd=str(mkdocs_dir),
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            creationflags=subprocess.CREATE_NEW_PROCESS_GROUP if os.name == "nt" else 0,
        )

        # Give it a moment to start
        import time

        time.sleep(2)

        # Check if it started successfully
        if _mkdocs_process.poll() is not None:
            # Process ended immediately, something went wrong
            stdout, stderr = _mkdocs_process.communicate()
            raise HTTPException(status_code=500, detail=f"Failed to start MkDocs: {stderr.decode()}")

        return {
            "status": "started",
            "message": f"MkDocs server started on port {MKDOCS_PORT}",
            "url": f"http://{MKDOCS_HOST}:{MKDOCS_PORT}",
            "pid": _mkdocs_process.pid,
        }

    except FileNotFoundError:
        raise HTTPException(
            status_code=500, detail="MkDocs not installed. Install with: pip install mkdocs mkdocs-material"
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to start MkDocs: {str(e)}")


@router.post("/api/mkdocs/stop")
def stop_mkdocs() -> Dict[str, str]:
    """Stop the MkDocs development server."""
    global _mkdocs_process

    if not is_mkdocs_running():
        return {"status": "not_running", "message": "MkDocs server is not running"}

    try:
        pid = get_mkdocs_pid()

        if pid:
            process = psutil.Process(pid)

            # Terminate gracefully
            if os.name == "nt":
                # Windows: use CTRL_BREAK_EVENT
                process.send_signal(signal.CTRL_BREAK_EVENT)
            else:
                # Unix: use SIGTERM
                process.terminate()

            # Wait for process to end
            try:
                process.wait(timeout=5)
            except psutil.TimeoutExpired:
                # Force kill if it doesn't stop gracefully
                process.kill()

        _mkdocs_process = None

        return {"status": "stopped", "message": "MkDocs server stopped successfully"}

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to stop MkDocs: {str(e)}")


@router.post("/api/mkdocs/restart")
def restart_mkdocs() -> Dict[str, Any]:
    """Restart the MkDocs development server."""
    if is_mkdocs_running():
        stop_mkdocs()

    import time

    time.sleep(1)

    return start_mkdocs()


@router.post("/api/mkdocs/build")
def build_mkdocs() -> Dict[str, str]:
    """Build the static MkDocs site."""
    mkdocs_dir = get_mkdocs_dir()

    try:
        # Run mkdocs build
        result = subprocess.run(
            ["mkdocs", "build", "--clean"], cwd=str(mkdocs_dir), capture_output=True, text=True, timeout=60
        )

        if result.returncode != 0:
            raise HTTPException(status_code=500, detail=f"Build failed: {result.stderr}")

        site_dir = mkdocs_dir / "site"

        return {
            "status": "success",
            "message": "MkDocs site built successfully",
            "output_dir": str(site_dir),
            "stdout": result.stdout,
        }

    except subprocess.TimeoutExpired:
        raise HTTPException(status_code=500, detail="Build timed out after 60 seconds")
    except FileNotFoundError:
        raise HTTPException(
            status_code=500, detail="MkDocs not installed. Install with: pip install mkdocs mkdocs-material"
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to build MkDocs: {str(e)}")


def get_documents_dir() -> Path:
    """Get the documents storage directory."""
    base = Path(__file__).parent.parent.parent / "data" / "documents"
    return base


def get_published_docs_dir() -> Path:
    """Get the MkDocs published documents directory."""
    base = Path(__file__).parent.parent.parent / "docs" / "generated" / "published"
    return base


def parse_frontmatter(content: str) -> Dict[str, Any]:
    """Parse YAML frontmatter from markdown content."""
    frontmatter = {}
    # Normalize line endings
    content = content.replace("\r\n", "\n")
    if content.startswith("---\n"):
        try:
            end = content.find("\n---\n", 4)
            if end != -1:
                import yaml

                frontmatter_text = content[4:end]
                frontmatter = yaml.safe_load(frontmatter_text) or {}
        except Exception:
            pass
    return frontmatter


@router.get("/api/mkdocs/documents/tree")
def get_documents_tree() -> Dict[str, Any]:
    """Get hierarchical tree structure of all published documents."""
    published_dir = get_published_docs_dir()

    if not published_dir.exists():
        return {"documents": [], "groups": {}}

    documents = []
    groups = defaultdict(list)

    for md_file in published_dir.glob("*.md"):
        try:
            with open(md_file, "r", encoding="utf-8") as f:
                content = f.read()

            frontmatter = parse_frontmatter(content)

            doc_info = {
                "filename": md_file.name,
                "title": frontmatter.get("title", md_file.stem),
                "template_type": frontmatter.get("template_type", "unknown"),
                "module": frontmatter.get("module", "unknown"),
                "created": frontmatter.get("created", None),
                "updated": frontmatter.get("updated", None),
                "url": f"/generated/published/{md_file.name}",
                "size": md_file.stat().st_size,
                "modified": datetime.fromtimestamp(md_file.stat().st_mtime).isoformat(),
            }

            documents.append(doc_info)

            # Group by template type
            template_type = doc_info["template_type"]
            groups[f"type_{template_type}"].append(doc_info)

            # Group by module
            module = doc_info["module"]
            if module and module != "unknown":
                groups[f"module_{module}"].append(doc_info)

        except Exception as e:
            print(f"Error processing {md_file}: {e}")
            continue

    # Sort documents by title
    documents.sort(key=lambda x: x["title"].lower())

    # Sort groups
    for key in groups:
        groups[key].sort(key=lambda x: x["title"].lower())

    return {
        "documents": documents,
        "groups": dict(groups),
        "total": len(documents),
        "by_type": {k.replace("type_", ""): len(v) for k, v in groups.items() if k.startswith("type_")},
        "by_module": {k.replace("module_", ""): len(v) for k, v in groups.items() if k.startswith("module_")},
    }


@router.get("/api/mkdocs/documents/search")
def search_documents(q: str = "", template_type: str = None, module: str = None) -> List[Dict[str, Any]]:
    """Search published documents by query, template type, or module."""
    published_dir = get_published_docs_dir()

    if not published_dir.exists():
        return []

    results = []
    query_lower = q.lower()

    for md_file in published_dir.glob("*.md"):
        try:
            with open(md_file, "r", encoding="utf-8") as f:
                content = f.read()

            frontmatter = parse_frontmatter(content)

            # Apply filters
            if template_type and frontmatter.get("template_type") != template_type:
                continue

            if module and frontmatter.get("module") != module:
                continue

            # Search in title and content
            title = frontmatter.get("title", md_file.stem)
            if query_lower:
                if query_lower not in title.lower() and query_lower not in content.lower():
                    continue

            doc_info = {
                "filename": md_file.name,
                "title": title,
                "template_type": frontmatter.get("template_type", "unknown"),
                "module": frontmatter.get("module", "unknown"),
                "created": frontmatter.get("created", None),
                "updated": frontmatter.get("updated", None),
                "url": f"/generated/published/{md_file.name}",
                "excerpt": content[:200].replace("---\n", "").strip()[:150] + "...",
            }

            results.append(doc_info)

        except Exception as e:
            print(f"Error searching {md_file}: {e}")
            continue

    # Sort by relevance (title matches first)
    if query_lower:
        results.sort(key=lambda x: (query_lower not in x["title"].lower(), x["title"].lower()))
    else:
        results.sort(key=lambda x: x["title"].lower())

    return results


@router.get("/api/mkdocs/documents/recent")
def get_recent_documents(limit: int = 10) -> List[Dict[str, Any]]:
    """Get recently updated documents."""
    published_dir = get_published_docs_dir()

    if not published_dir.exists():
        return []

    documents = []

    for md_file in published_dir.glob("*.md"):
        try:
            with open(md_file, "r", encoding="utf-8") as f:
                content = f.read()

            frontmatter = parse_frontmatter(content)

            doc_info = {
                "filename": md_file.name,
                "title": frontmatter.get("title", md_file.stem),
                "template_type": frontmatter.get("template_type", "unknown"),
                "module": frontmatter.get("module", "unknown"),
                "created": frontmatter.get("created", None),
                "updated": frontmatter.get("updated", None),
                "url": f"/generated/published/{md_file.name}",
                "modified": md_file.stat().st_mtime,
            }

            documents.append(doc_info)

        except Exception as e:
            print(f"Error processing {md_file}: {e}")
            continue

    # Sort by modification time (most recent first)
    documents.sort(key=lambda x: x["modified"], reverse=True)

    # Remove the modified timestamp from results
    for doc in documents[:limit]:
        doc["modified"] = datetime.fromtimestamp(doc["modified"]).isoformat()

    return documents[:limit]


@router.get("/api/mkdocs/documents/{filename}/metadata")
def get_document_metadata(filename: str) -> Dict[str, Any]:
    """Get detailed metadata for a specific document."""
    published_dir = get_published_docs_dir()
    doc_path = published_dir / filename

    if not doc_path.exists():
        raise HTTPException(status_code=404, detail=f"Document '{filename}' not found")

    try:
        with open(doc_path, "r", encoding="utf-8") as f:
            content = f.read()

        frontmatter = parse_frontmatter(content)

        # Try to load source document for additional metadata
        docs_dir = get_documents_dir()
        source_metadata = {}

        # Find matching source document
        for doc_file in docs_dir.glob("*.json"):
            try:
                with open(doc_file, "r", encoding="utf-8") as f:
                    doc_data = json.load(f)
                    if doc_data.get("title") == frontmatter.get("title"):
                        source_metadata = {
                            "doc_id": doc_data.get("id"),
                            "atom_ids": doc_data.get("atom_ids", []),
                            "approval_status": doc_data.get("approval_status"),
                            "version": doc_data.get("version"),
                            "reviewed_by": doc_data.get("reviewed_by"),
                            "published_at": doc_data.get("published_at"),
                        }
                        break
            except Exception:
                continue

        stats = doc_path.stat()

        return {
            "filename": filename,
            "title": frontmatter.get("title", filename),
            "template_type": frontmatter.get("template_type", "unknown"),
            "module": frontmatter.get("module", "unknown"),
            "created": frontmatter.get("created", None),
            "updated": frontmatter.get("updated", None),
            "url": f"/generated/published/{filename}",
            "size": stats.st_size,
            "modified": datetime.fromtimestamp(stats.st_mtime).isoformat(),
            "word_count": len(content.split()),
            "line_count": content.count("\n"),
            **source_metadata,
        }

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to get metadata: {str(e)}")


# ============================================================================
# PHASE 2: SEARCH & DISCOVERY
# ============================================================================


@router.get("/api/mkdocs/documents/semantic-search")
async def semantic_search_documents(query: str, limit: int = 10) -> List[Dict[str, Any]]:
    """Semantic search across documents using RAG system."""
    try:
        from .rag import entity_rag

        # Use RAG to find relevant atoms
        rag_results = entity_rag(query, top_k=limit * 3)  # Get more to filter by documents

        # Get unique documents that contain these atoms
        published_dir = get_published_docs_dir()
        if not published_dir.exists():
            return []

        doc_scores = defaultdict(lambda: {"score": 0.0, "matched_atoms": []})

        # Load all documents and match atoms
        for md_file in published_dir.glob("*.md"):
            try:
                with open(md_file, "r", encoding="utf-8") as f:
                    content = f.read()

                frontmatter = parse_frontmatter(content)

                # Check if any RAG result atoms are in this document's content or metadata
                for rag_result in rag_results:
                    atom_id = rag_result.get("atom_id", "")
                    if atom_id in content or atom_id in str(frontmatter):
                        doc_scores[md_file.name]["score"] += rag_result.get("similarity", 0.0)
                        doc_scores[md_file.name]["matched_atoms"].append(
                            {
                                "atom_id": atom_id,
                                "name": rag_result.get("name", ""),
                                "relevance": rag_result.get("similarity", 0.0),
                            }
                        )

                # Also check content similarity directly
                if query.lower() in content.lower():
                    doc_scores[md_file.name]["score"] += 0.5

            except Exception as e:
                print(f"Error processing {md_file}: {e}")
                continue

        # Convert to list and add document info
        results = []
        for filename, score_data in doc_scores.items():
            try:
                doc_path = published_dir / filename
                with open(doc_path, "r", encoding="utf-8") as f:
                    content = f.read()

                frontmatter = parse_frontmatter(content)

                results.append(
                    {
                        "filename": filename,
                        "title": frontmatter.get("title", filename),
                        "template_type": frontmatter.get("template_type", "unknown"),
                        "module": frontmatter.get("module", "unknown"),
                        "url": f"/generated/published/{filename}",
                        "relevance_score": score_data["score"],
                        "matched_atoms": score_data["matched_atoms"][:5],  # Top 5 atoms
                        "excerpt": content[:300].replace("---\n", "").strip()[:200] + "...",
                    }
                )
            except Exception:
                continue

        # Sort by relevance score and return top results
        results.sort(key=lambda x: x["relevance_score"], reverse=True)
        return results[:limit]

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Semantic search failed: {str(e)}")


@router.get("/api/mkdocs/documents/{filename}/related")
async def get_related_documents(filename: str, limit: int = 5) -> List[Dict[str, Any]]:
    """Find related documents based on shared atoms and content similarity."""
    published_dir = get_published_docs_dir()
    doc_path = published_dir / filename

    if not doc_path.exists():
        raise HTTPException(status_code=404, detail=f"Document '{filename}' not found")

    try:
        # Load source document
        with open(doc_path, "r", encoding="utf-8") as f:
            source_content = f.read()

        source_frontmatter = parse_frontmatter(source_content)
        source_module = source_frontmatter.get("module")
        source_type = source_frontmatter.get("template_type")

        # Extract atom IDs from source document
        source_atom_ids = set()
        # Look for atom IDs in content (format: atom-xxx-yyy)
        import re

        atom_pattern = r"atom-[\w-]+"
        source_atom_ids = set(re.findall(atom_pattern, source_content))

        # Calculate similarity with other documents
        related_docs = []

        for md_file in published_dir.glob("*.md"):
            if md_file.name == filename:
                continue

            try:
                with open(md_file, "r", encoding="utf-8") as f:
                    content = f.read()

                frontmatter = parse_frontmatter(content)

                # Find shared atoms
                doc_atom_ids = set(re.findall(atom_pattern, content))
                shared_atoms = source_atom_ids.intersection(doc_atom_ids)

                # Calculate similarity score
                score = 0.0

                # Shared atoms (highest weight)
                if shared_atoms:
                    score += len(shared_atoms) * 2.0

                # Same module (medium weight)
                if frontmatter.get("module") == source_module and source_module != "unknown":
                    score += 1.5

                # Same template type (low weight)
                if frontmatter.get("template_type") == source_type and source_type != "unknown":
                    score += 0.5

                if score > 0:
                    related_docs.append(
                        {
                            "filename": md_file.name,
                            "title": frontmatter.get("title", md_file.stem),
                            "template_type": frontmatter.get("template_type", "unknown"),
                            "module": frontmatter.get("module", "unknown"),
                            "url": f"/generated/published/{md_file.name}",
                            "similarity_score": score,
                            "shared_atoms_count": len(shared_atoms),
                            "shared_atoms": list(shared_atoms)[:10],  # Top 10
                            "relation_type": (
                                "same_module" if frontmatter.get("module") == source_module else "shared_atoms"
                            ),
                        }
                    )

            except Exception as e:
                print(f"Error processing {md_file}: {e}")
                continue

        # Sort by similarity score
        related_docs.sort(key=lambda x: x["similarity_score"], reverse=True)
        return related_docs[:limit]

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to find related documents: {str(e)}")


@router.get("/api/mkdocs/documents/{filename}/ai-recommendations")
async def get_ai_recommendations(filename: str, limit: int = 5) -> Dict[str, Any]:
    """Get AI-powered document recommendations using RAG."""
    published_dir = get_published_docs_dir()
    doc_path = published_dir / filename

    if not doc_path.exists():
        raise HTTPException(status_code=404, detail=f"Document '{filename}' not found")

    try:
        from .rag import entity_rag, path_rag

        # Load source document
        with open(doc_path, "r", encoding="utf-8") as f:
            content = f.read()

        frontmatter = parse_frontmatter(content)

        # Extract key concepts from document for RAG query
        title = frontmatter.get("title", "")
        module = frontmatter.get("module", "")

        # Create intelligent query
        query = f"{title} {module}"

        # Use RAG to find related atoms
        entity_results = entity_rag(query, top_k=10)
        path_results = path_rag(query, top_k=10)

        # Combine and deduplicate atom IDs
        all_atom_ids = set()
        for result in entity_results + path_results:
            all_atom_ids.add(result.get("atom_id", ""))

        # Find documents containing these atoms
        recommendations = []

        for md_file in published_dir.glob("*.md"):
            if md_file.name == filename:
                continue

            try:
                with open(md_file, "r", encoding="utf-8") as f:
                    doc_content = f.read()

                doc_frontmatter = parse_frontmatter(doc_content)

                # Count matching atoms
                matches = sum(1 for atom_id in all_atom_ids if atom_id in doc_content)

                if matches > 0:
                    recommendations.append(
                        {
                            "filename": md_file.name,
                            "title": doc_frontmatter.get("title", md_file.stem),
                            "template_type": doc_frontmatter.get("template_type", "unknown"),
                            "module": doc_frontmatter.get("module", "unknown"),
                            "url": f"/generated/published/{md_file.name}",
                            "ai_score": matches,
                            "reason": f"Contains {matches} atoms related to this document",
                        }
                    )

            except Exception:
                continue

        # Sort by AI score
        recommendations.sort(key=lambda x: x["ai_score"], reverse=True)

        return {
            "source_document": {"filename": filename, "title": title, "module": module},
            "recommendations": recommendations[:limit],
            "total_found": len(recommendations),
            "rag_atoms_analyzed": len(all_atom_ids),
        }

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"AI recommendations failed: {str(e)}")


# ============================================================================
# PHASE 3: COLLABORATION FEATURES
# ============================================================================


@router.get("/api/mkdocs/documents/{filename}/export/pdf")
async def export_document_to_pdf(filename: str):
    """Export a document to PDF format."""
    try:
        from io import BytesIO

        import markdown
        from bs4 import BeautifulSoup
        from reportlab.lib.enums import TA_CENTER, TA_LEFT
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import inch
        from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer

        published_dir = get_published_docs_dir()
        doc_path = published_dir / filename

        if not doc_path.exists():
            raise HTTPException(status_code=404, detail="Document not found")

        # Read document content
        with open(doc_path, "r", encoding="utf-8") as f:
            content = f.read()

        # Parse frontmatter
        frontmatter = {}
        if content.startswith("---"):
            parts = content.split("---", 2)
            if len(parts) >= 3:
                import yaml

                frontmatter = yaml.safe_load(parts[1]) or {}
                content = parts[2]

        # Convert markdown to HTML then to plain text for PDF
        html = markdown.markdown(content)
        soup = BeautifulSoup(html, "html.parser")

        # Create PDF
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter, rightMargin=72, leftMargin=72, topMargin=72, bottomMargin=18)

        # Styles
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            "CustomTitle",
            parent=styles["Heading1"],
            fontSize=24,
            textColor="#1a1a1a",
            spaceAfter=30,
            alignment=TA_CENTER,
        )
        heading_style = ParagraphStyle(
            "CustomHeading", parent=styles["Heading2"], fontSize=16, textColor="#333333", spaceAfter=12, spaceBefore=12
        )
        body_style = ParagraphStyle(
            "CustomBody", parent=styles["BodyText"], fontSize=11, textColor="#444444", spaceAfter=12, alignment=TA_LEFT
        )

        # Build PDF content
        story = []

        # Add title
        title = frontmatter.get("title", filename.replace(".md", ""))
        story.append(Paragraph(title, title_style))
        story.append(Spacer(1, 12))

        # Add metadata
        if frontmatter:
            metadata_lines = []
            if "template_type" in frontmatter:
                metadata_lines.append(f"<b>Type:</b> {frontmatter['template_type']}")
            if "module" in frontmatter:
                metadata_lines.append(f"<b>Module:</b> {frontmatter['module']}")
            if "owner" in frontmatter:
                metadata_lines.append(f"<b>Owner:</b> {frontmatter['owner']}")
            if "version" in frontmatter:
                metadata_lines.append(f"<b>Version:</b> {frontmatter['version']}")

            if metadata_lines:
                story.append(Paragraph(" | ".join(metadata_lines), body_style))
                story.append(Spacer(1, 20))

        # Add content
        for element in soup.find_all(["h1", "h2", "h3", "h4", "p", "ul", "ol"]):
            if element.name in ["h1", "h2", "h3", "h4"]:
                story.append(Paragraph(element.get_text(), heading_style))
            elif element.name == "p":
                story.append(Paragraph(element.get_text(), body_style))
            elif element.name in ["ul", "ol"]:
                for li in element.find_all("li"):
                    story.append(Paragraph(f"• {li.get_text()}", body_style))

        # Build PDF
        doc.build(story)
        buffer.seek(0)

        from fastapi.responses import StreamingResponse

        return StreamingResponse(
            buffer,
            media_type="application/pdf",
            headers={"Content-Disposition": f"attachment; filename={filename.replace('.md', '.pdf')}"},
        )

    except ImportError:
        raise HTTPException(
            status_code=500, detail="PDF export requires reportlab library. Install with: pip install reportlab"
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"PDF export failed: {str(e)}")


@router.get("/api/mkdocs/documents/{filename}/export/docx")
async def export_document_to_docx(filename: str):
    """Export a document to Word (DOCX) format."""
    try:
        from io import BytesIO

        import markdown
        from bs4 import BeautifulSoup
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Inches, Pt, RGBColor

        published_dir = get_published_docs_dir()
        doc_path = published_dir / filename

        if not doc_path.exists():
            raise HTTPException(status_code=404, detail="Document not found")

        # Read document content
        with open(doc_path, "r", encoding="utf-8") as f:
            content = f.read()

        # Parse frontmatter
        frontmatter = {}
        if content.startswith("---"):
            parts = content.split("---", 2)
            if len(parts) >= 3:
                import yaml

                frontmatter = yaml.safe_load(parts[1]) or {}
                content = parts[2]

        # Convert markdown to HTML
        html = markdown.markdown(content)
        soup = BeautifulSoup(html, "html.parser")

        # Create Word document
        doc = Document()

        # Add title
        title = frontmatter.get("title", filename.replace(".md", ""))
        title_para = doc.add_heading(title, level=0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add metadata table
        if frontmatter:
            doc.add_paragraph()  # Spacer
            metadata_table = doc.add_table(rows=0, cols=2)
            metadata_table.style = "Light Grid Accent 1"

            if "template_type" in frontmatter:
                row = metadata_table.add_row()
                row.cells[0].text = "Type"
                row.cells[1].text = str(frontmatter["template_type"])

            if "module" in frontmatter:
                row = metadata_table.add_row()
                row.cells[0].text = "Module"
                row.cells[1].text = str(frontmatter["module"])

            if "owner" in frontmatter:
                row = metadata_table.add_row()
                row.cells[0].text = "Owner"
                row.cells[1].text = str(frontmatter["owner"])

            if "steward" in frontmatter:
                row = metadata_table.add_row()
                row.cells[0].text = "Steward"
                row.cells[1].text = str(frontmatter["steward"])

            if "version" in frontmatter:
                row = metadata_table.add_row()
                row.cells[0].text = "Version"
                row.cells[1].text = str(frontmatter["version"])

            if "status" in frontmatter:
                row = metadata_table.add_row()
                row.cells[0].text = "Status"
                row.cells[1].text = str(frontmatter["status"])

            doc.add_paragraph()  # Spacer

        # Add content
        for element in soup.find_all(["h1", "h2", "h3", "h4", "p", "ul", "ol"]):
            if element.name == "h1":
                doc.add_heading(element.get_text(), level=1)
            elif element.name == "h2":
                doc.add_heading(element.get_text(), level=2)
            elif element.name == "h3":
                doc.add_heading(element.get_text(), level=3)
            elif element.name == "h4":
                doc.add_heading(element.get_text(), level=4)
            elif element.name == "p":
                doc.add_paragraph(element.get_text())
            elif element.name in ["ul", "ol"]:
                for li in element.find_all("li"):
                    doc.add_paragraph(li.get_text(), style="List Bullet")

        # Save to buffer
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        from fastapi.responses import StreamingResponse

        return StreamingResponse(
            buffer,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename={filename.replace('.md', '.docx')}"},
        )

    except ImportError:
        raise HTTPException(
            status_code=500, detail="DOCX export requires python-docx library. Install with: pip install python-docx"
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"DOCX export failed: {str(e)}")


# ============================================================================
# PHASE 4: ANALYTICS & INSIGHTS
# ============================================================================

# In-memory analytics store (in production, use a database)
analytics_store = {"document_views": defaultdict(int), "search_queries": [], "view_history": []}


@router.post("/api/mkdocs/analytics/track-view")
async def track_document_view(request: dict):
    """Track a document view for analytics."""
    try:
        filename = request.get("filename")
        timestamp = request.get("timestamp", None)

        if not filename:
            raise HTTPException(status_code=400, detail="Filename is required")

        # Increment view count
        analytics_store["document_views"][filename] += 1

        # Add to view history
        analytics_store["view_history"].append(
            {
                "filename": filename,
                "timestamp": timestamp or str(datetime.now()),
                "user": "anonymous",  # In production, use actual user ID
            }
        )

        # Keep only last 1000 views
        if len(analytics_store["view_history"]) > 1000:
            analytics_store["view_history"] = analytics_store["view_history"][-1000:]

        return {"status": "success", "views": analytics_store["document_views"][filename]}

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to track view: {str(e)}")


@router.post("/api/mkdocs/analytics/track-search")
async def track_search_query(request: dict):
    """Track a search query for analytics."""
    try:
        query = request.get("query")
        mode = request.get("mode", "basic")
        results_count = request.get("results_count", 0)

        if not query:
            return {"status": "ignored"}

        analytics_store["search_queries"].append(
            {"query": query, "mode": mode, "results_count": results_count, "timestamp": str(datetime.now())}
        )

        # Keep only last 500 queries
        if len(analytics_store["search_queries"]) > 500:
            analytics_store["search_queries"] = analytics_store["search_queries"][-500:]

        return {"status": "success"}

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to track search: {str(e)}")


@router.get("/api/mkdocs/analytics/usage-stats")
async def get_usage_statistics():
    """Get overall usage statistics."""
    try:
        # Calculate stats
        total_views = sum(analytics_store["document_views"].values())
        total_documents = len(analytics_store["document_views"])
        total_searches = len(analytics_store["search_queries"])

        # Top 10 most viewed documents
        top_documents = sorted(analytics_store["document_views"].items(), key=lambda x: x[1], reverse=True)[:10]

        # Recent activity (last 24 hours)
        now = datetime.now()
        one_day_ago = now - timedelta(days=1)
        recent_views = (
            [
                v
                for v in analytics_store["view_history"]
                if datetime.fromisoformat(v["timestamp"].replace("Z", "+00:00")) > one_day_ago
            ]
            if analytics_store["view_history"]
            else []
        )

        # Top search queries
        query_counts = defaultdict(int)
        for search in analytics_store["search_queries"]:
            query_counts[search["query"]] += 1

        top_searches = sorted(query_counts.items(), key=lambda x: x[1], reverse=True)[:10]

        return {
            "total_views": total_views,
            "total_documents_viewed": total_documents,
            "total_searches": total_searches,
            "top_documents": [{"filename": filename, "views": views} for filename, views in top_documents],
            "recent_activity_count": len(recent_views),
            "top_searches": [{"query": query, "count": count} for query, count in top_searches],
        }

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to get usage stats: {str(e)}")


@router.get("/api/mkdocs/analytics/document-health")
async def get_document_health_metrics():
    """Calculate document health metrics (freshness, completeness, quality)."""
    try:
        published_dir = get_published_docs_dir()
        if not published_dir.exists():
            return []

        health_metrics = []
        now = datetime.now()

        for md_file in published_dir.glob("*.md"):
            try:
                with open(md_file, "r", encoding="utf-8") as f:
                    content = f.read()

                # Parse frontmatter
                frontmatter = {}
                if content.startswith("---"):
                    parts = content.split("---", 2)
                    if len(parts) >= 3:
                        import yaml

                        frontmatter = yaml.safe_load(parts[1]) or {}
                        content = parts[2]

                # Calculate metrics
                word_count = len(content.split())
                has_owner = bool(frontmatter.get("owner"))
                has_version = bool(frontmatter.get("version"))
                has_status = bool(frontmatter.get("status"))

                # Freshness score (based on modification time)
                mod_time = datetime.fromtimestamp(md_file.stat().st_mtime)
                days_old = (now - mod_time).days
                freshness_score = max(0, 100 - (days_old * 2))  # Lose 2 points per day

                # Completeness score (based on metadata presence)
                completeness_score = 0
                if has_owner:
                    completeness_score += 25
                if has_version:
                    completeness_score += 25
                if has_status:
                    completeness_score += 25
                if word_count > 100:
                    completeness_score += 25

                # Quality score (composite)
                quality_score = (freshness_score * 0.4) + (completeness_score * 0.6)

                # Health status
                if quality_score >= 80:
                    health_status = "excellent"
                elif quality_score >= 60:
                    health_status = "good"
                elif quality_score >= 40:
                    health_status = "fair"
                else:
                    health_status = "poor"

                health_metrics.append(
                    {
                        "filename": md_file.name,
                        "title": frontmatter.get("title", md_file.name.replace(".md", "")),
                        "freshness_score": round(freshness_score, 1),
                        "completeness_score": completeness_score,
                        "quality_score": round(quality_score, 1),
                        "health_status": health_status,
                        "days_since_update": days_old,
                        "word_count": word_count,
                        "has_owner": has_owner,
                        "views": analytics_store["document_views"].get(md_file.name, 0),
                    }
                )

            except Exception as e:
                print(f"Error processing {md_file.name}: {e}")
                continue

        # Sort by quality score descending
        health_metrics.sort(key=lambda x: x["quality_score"], reverse=True)

        return health_metrics

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to calculate document health: {str(e)}")


# ============================================================================
# ENHANCED FEATURES: VERSIONING & COMPARISON
# ============================================================================

# Document version history store (in production, use a database)
version_store = defaultdict(list)


@router.post("/api/mkdocs/documents/{filename}/save-version")
async def save_document_version(filename: str, request: dict):
    """Save a version snapshot of a document."""
    try:
        published_dir = get_published_docs_dir()
        doc_path = published_dir / filename

        if not doc_path.exists():
            raise HTTPException(status_code=404, detail="Document not found")

        # Read current content
        with open(doc_path, "r", encoding="utf-8") as f:
            content = f.read()

        frontmatter = parse_frontmatter(content)

        # Create version entry
        version = {
            "version_id": len(version_store[filename]) + 1,
            "filename": filename,
            "content": content,
            "frontmatter": frontmatter,
            "timestamp": datetime.now().isoformat(),
            "author": request.get("author", "anonymous"),
            "comment": request.get("comment", ""),
            "word_count": len(content.split()),
            "hash": hash(content),
        }

        version_store[filename].append(version)

        # Keep only last 50 versions per document
        if len(version_store[filename]) > 50:
            version_store[filename] = version_store[filename][-50:]

        return {"status": "success", "version_id": version["version_id"], "timestamp": version["timestamp"]}

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to save version: {str(e)}")


@router.get("/api/mkdocs/documents/{filename}/versions")
async def get_document_versions(filename: str, limit: int = 10):
    """Get version history for a document."""
    try:
        versions = version_store.get(filename, [])

        # Return metadata only (not full content)
        version_list = []
        for v in reversed(versions[-limit:]):
            version_list.append(
                {
                    "version_id": v["version_id"],
                    "timestamp": v["timestamp"],
                    "author": v["author"],
                    "comment": v["comment"],
                    "word_count": v["word_count"],
                }
            )

        return {"filename": filename, "total_versions": len(versions), "versions": version_list}

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to get versions: {str(e)}")


@router.get("/api/mkdocs/documents/{filename}/compare")
async def compare_document_versions(filename: str, version1: int = None, version2: int = None):
    """Compare two versions of a document or current version with a previous one."""
    try:
        published_dir = get_published_docs_dir()
        doc_path = published_dir / filename

        if not doc_path.exists():
            raise HTTPException(status_code=404, detail="Document not found")

        versions = version_store.get(filename, [])

        # Get current content
        with open(doc_path, "r", encoding="utf-8") as f:
            current_content = f.read()

        # Determine what to compare
        if version2 is None:
            # Compare current with a version
            content1 = current_content
            content2 = versions[version1 - 1]["content"] if version1 and version1 <= len(versions) else ""
            label1 = "Current"
            label2 = f"Version {version1}"
        else:
            # Compare two versions
            content1 = versions[version1 - 1]["content"] if version1 and version1 <= len(versions) else ""
            content2 = versions[version2 - 1]["content"] if version2 and version2 <= len(versions) else ""
            label1 = f"Version {version1}"
            label2 = f"Version {version2}"

        # Simple diff calculation
        import difflib

        diff = difflib.unified_diff(
            content2.splitlines(keepends=True),
            content1.splitlines(keepends=True),
            fromfile=label2,
            tofile=label1,
            lineterm="",
        )

        diff_text = "".join(diff)

        # Calculate statistics
        lines1 = content1.count("\n")
        lines2 = content2.count("\n")
        words1 = len(content1.split())
        words2 = len(content2.split())

        return {
            "filename": filename,
            "comparison": {
                "label1": label1,
                "label2": label2,
                "diff": diff_text,
                "stats": {
                    "lines_added": max(0, lines1 - lines2),
                    "lines_removed": max(0, lines2 - lines1),
                    "words_added": max(0, words1 - words2),
                    "words_removed": max(0, words2 - words1),
                },
            },
        }

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to compare versions: {str(e)}")


# ============================================================================
# ENHANCED FEATURES: DOCUMENT TEMPLATES
# ============================================================================


@router.get("/api/mkdocs/templates")
async def get_document_templates():
    """Get available document templates."""
    templates = [
        {
            "id": "sop",
            "name": "Standard Operating Procedure",
            "description": "Template for creating SOPs with steps and compliance notes",
            "template_type": "sop",
            "sections": ["Purpose", "Scope", "Procedure", "References", "Approval"],
        },
        {
            "id": "technical_design",
            "name": "Technical Design Document",
            "description": "Template for technical specifications and architecture",
            "template_type": "technical_design",
            "sections": ["Overview", "Architecture", "Components", "Data Flow", "Security", "Testing"],
        },
        {
            "id": "process_map",
            "name": "Process Map",
            "description": "Template for documenting business processes",
            "template_type": "process_map",
            "sections": ["Process Overview", "Inputs", "Steps", "Outputs", "Metrics"],
        },
        {
            "id": "document_inventory",
            "name": "Document Inventory",
            "description": "Template for cataloging document collections",
            "template_type": "document_inventory",
            "sections": ["Overview", "Document List", "Ownership", "Review Schedule"],
        },
        {
            "id": "api_documentation",
            "name": "API Documentation",
            "description": "Template for API endpoint documentation",
            "template_type": "api_documentation",
            "sections": ["Endpoint", "Authentication", "Request", "Response", "Examples", "Error Codes"],
        },
    ]

    return {"templates": templates, "total": len(templates)}


@router.post("/api/mkdocs/documents/create-from-template")
async def create_document_from_template(request: dict):
    """Create a new document from a template."""
    try:
        template_id = request.get("template_id")
        title = request.get("title")
        module = request.get("module", "general")
        owner = request.get("owner", "unassigned")

        if not template_id or not title:
            raise HTTPException(status_code=400, detail="template_id and title are required")

        # Get template
        templates_response = await get_document_templates()
        templates = {t["id"]: t for t in templates_response["templates"]}

        if template_id not in templates:
            raise HTTPException(status_code=404, detail="Template not found")

        template = templates[template_id]

        # Generate document content
        frontmatter = {
            "title": title,
            "template_type": template["template_type"],
            "module": module,
            "owner": owner,
            "steward": owner,
            "created": datetime.now().isoformat(),
            "updated": datetime.now().isoformat(),
            "version": "1.0.0",
            "status": "draft",
        }

        # Build markdown content
        import yaml

        content = "---\n"
        content += yaml.dump(frontmatter, default_flow_style=False)
        content += "---\n\n"
        content += f"# {title}\n\n"

        # Add template sections
        for section in template["sections"]:
            content += f"## {section}\n\n"
            content += f"*TODO: Complete the {section} section*\n\n"

        # Save to published directory
        published_dir = get_published_docs_dir()
        published_dir.mkdir(parents=True, exist_ok=True)

        # Generate filename
        safe_title = title.lower().replace(" ", "_").replace("/", "_")
        filename = f"{template['template_type']}_{safe_title}.md"
        filepath = published_dir / filename

        # Check if file exists
        if filepath.exists():
            raise HTTPException(status_code=409, detail="Document already exists")

        with open(filepath, "w", encoding="utf-8") as f:
            f.write(content)

        return {
            "status": "success",
            "filename": filename,
            "url": f"/generated/published/{filename}",
            "message": f"Document created from {template['name']} template",
        }

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to create document: {str(e)}")


# ============================================================================
# ENHANCED FEATURES: DEPENDENCY GRAPH
# ============================================================================


@router.get("/api/mkdocs/documents/{filename}/dependencies")
async def get_document_dependencies(filename: str):
    """Get dependency graph for a document based on atom relationships."""
    try:
        published_dir = get_published_docs_dir()
        doc_path = published_dir / filename

        if not doc_path.exists():
            raise HTTPException(status_code=404, detail="Document not found")

        # Read document and extract atoms
        with open(doc_path, "r", encoding="utf-8") as f:
            content = f.read()

        frontmatter = parse_frontmatter(content)

        # Extract atom IDs
        atom_pattern = r"atom-[\w-]+"
        doc_atoms = set(re.findall(atom_pattern, content))

        # Find all documents and their atoms
        doc_atoms_map = {}
        for md_file in published_dir.glob("*.md"):
            try:
                with open(md_file, "r", encoding="utf-8") as f:
                    file_content = f.read()
                file_atoms = set(re.findall(atom_pattern, file_content))
                if file_atoms:
                    doc_atoms_map[md_file.name] = file_atoms
            except Exception:
                continue

        # Build dependency graph
        nodes = [{"id": filename, "label": frontmatter.get("title", filename), "type": "source"}]
        edges = []

        for other_doc, other_atoms in doc_atoms_map.items():
            if other_doc == filename:
                continue

            # Find shared atoms
            shared = doc_atoms.intersection(other_atoms)
            if shared:
                # Determine direction
                if len(shared) >= 3:  # Strong dependency
                    edge_type = "strong"
                else:
                    edge_type = "weak"

                edges.append(
                    {
                        "source": filename,
                        "target": other_doc,
                        "type": edge_type,
                        "shared_atoms": len(shared),
                        "atoms": list(shared)[:5],
                    }
                )

                # Add node if not already present
                if not any(n["id"] == other_doc for n in nodes):
                    try:
                        with open(published_dir / other_doc, "r", encoding="utf-8") as f:
                            other_content = f.read()
                        other_frontmatter = parse_frontmatter(other_content)
                        nodes.append(
                            {"id": other_doc, "label": other_frontmatter.get("title", other_doc), "type": "dependency"}
                        )
                    except Exception:
                        nodes.append({"id": other_doc, "label": other_doc, "type": "dependency"})

        return {
            "source_document": filename,
            "atoms_count": len(doc_atoms),
            "dependencies_count": len(edges),
            "graph": {"nodes": nodes, "edges": edges},
        }

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to get dependencies: {str(e)}")


# ============================================================================
# ENHANCED FEATURES: ADVANCED SEARCH WITH HIGHLIGHTING
# ============================================================================


@router.get("/api/mkdocs/documents/advanced-search")
async def advanced_search(
    query: str,
    template_type: str = None,
    module: str = None,
    owner: str = None,
    min_score: float = 0.0,
    limit: int = 20,
):
    """Advanced search with content highlighting and relevance scoring."""
    try:
        published_dir = get_published_docs_dir()
        if not published_dir.exists():
            return []

        results = []
        query_terms = query.lower().split()

        for md_file in published_dir.glob("*.md"):
            try:
                with open(md_file, "r", encoding="utf-8") as f:
                    content = f.read()

                frontmatter = parse_frontmatter(content)

                # Apply filters
                if template_type and frontmatter.get("template_type") != template_type:
                    continue
                if module and frontmatter.get("module") != module:
                    continue
                if owner and frontmatter.get("owner") != owner:
                    continue

                # Calculate relevance score
                title = frontmatter.get("title", md_file.stem)
                content_lower = content.lower()
                title_lower = title.lower()

                score = 0.0

                # Title matches (highest weight: 10 points per term)
                for term in query_terms:
                    if term in title_lower:
                        score += 10.0

                # Content matches (1 point per occurrence, max 20)
                for term in query_terms:
                    occurrences = content_lower.count(term)
                    score += min(occurrences, 20)

                # Exact phrase bonus (5 points)
                if query.lower() in content_lower:
                    score += 5.0

                if score < min_score:
                    continue

                # Extract highlighted snippets
                snippets = []
                lines = content.split("\n")
                for i, line in enumerate(lines):
                    line_lower = line.lower()
                    if any(term in line_lower for term in query_terms):
                        # Add context (1 line before and after)
                        context_start = max(0, i - 1)
                        context_end = min(len(lines), i + 2)
                        context_lines = lines[context_start:context_end]

                        # Highlight matches
                        highlighted = []
                        for ctx_line in context_lines:
                            for term in query_terms:
                                # Case-insensitive replace with highlighting
                                pattern = re.compile(re.escape(term), re.IGNORECASE)
                                ctx_line = pattern.sub(f"<mark>{term}</mark>", ctx_line)
                            highlighted.append(ctx_line)

                        snippet = "\n".join(highlighted)
                        if snippet not in [s["text"] for s in snippets]:
                            snippets.append({"text": snippet, "line_number": i + 1})

                        if len(snippets) >= 3:  # Limit snippets
                            break

                results.append(
                    {
                        "filename": md_file.name,
                        "title": title,
                        "template_type": frontmatter.get("template_type", "unknown"),
                        "module": frontmatter.get("module", "unknown"),
                        "owner": frontmatter.get("owner", "unknown"),
                        "url": f"/generated/published/{md_file.name}",
                        "relevance_score": round(score, 2),
                        "snippets": snippets,
                        "total_matches": sum(content_lower.count(term) for term in query_terms),
                    }
                )

            except Exception as e:
                print(f"Error searching {md_file}: {e}")
                continue

        # Sort by relevance score
        results.sort(key=lambda x: x["relevance_score"], reverse=True)

        return results[:limit]

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Advanced search failed: {str(e)}")


# ============================================================================
# ENHANCED FEATURES: BULK OPERATIONS
# ============================================================================


@router.post("/api/mkdocs/documents/bulk-update")
async def bulk_update_documents(request: dict):
    """Bulk update multiple documents."""
    try:
        filenames = request.get("filenames", [])
        updates = request.get("updates", {})

        if not filenames:
            raise HTTPException(status_code=400, detail="No filenames provided")

        published_dir = get_published_docs_dir()
        results = {"success": [], "failed": []}

        for filename in filenames:
            try:
                doc_path = published_dir / filename
                if not doc_path.exists():
                    results["failed"].append({"filename": filename, "reason": "File not found"})
                    continue

                # Read document
                with open(doc_path, "r", encoding="utf-8") as f:
                    content = f.read()

                # Parse and update frontmatter
                if content.startswith("---"):
                    parts = content.split("---", 2)
                    if len(parts) >= 3:
                        import yaml

                        frontmatter = yaml.safe_load(parts[1]) or {}

                        # Apply updates
                        for key, value in updates.items():
                            frontmatter[key] = value

                        # Rebuild content
                        new_content = "---\n"
                        new_content += yaml.dump(frontmatter, default_flow_style=False)
                        new_content += "---"
                        new_content += parts[2]

                        # Save
                        with open(doc_path, "w", encoding="utf-8") as f:
                            f.write(new_content)

                        results["success"].append(filename)

            except Exception as e:
                results["failed"].append({"filename": filename, "reason": str(e)})

        return {
            "status": "completed",
            "total": len(filenames),
            "success_count": len(results["success"]),
            "failed_count": len(results["failed"]),
            "results": results,
        }

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Bulk update failed: {str(e)}")


@router.post("/api/mkdocs/documents/bulk-delete")
async def bulk_delete_documents(request: dict):
    """Bulk delete multiple documents."""
    try:
        filenames = request.get("filenames", [])
        confirm = request.get("confirm", False)

        if not confirm:
            raise HTTPException(status_code=400, detail="Deletion not confirmed. Set 'confirm': true")

        if not filenames:
            raise HTTPException(status_code=400, detail="No filenames provided")

        published_dir = get_published_docs_dir()
        results = {"deleted": [], "failed": []}

        for filename in filenames:
            try:
                doc_path = published_dir / filename
                if doc_path.exists():
                    doc_path.unlink()
                    results["deleted"].append(filename)
                else:
                    results["failed"].append({"filename": filename, "reason": "File not found"})

            except Exception as e:
                results["failed"].append({"filename": filename, "reason": str(e)})

        return {
            "status": "completed",
            "total": len(filenames),
            "deleted_count": len(results["deleted"]),
            "failed_count": len(results["failed"]),
            "results": results,
        }

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Bulk delete failed: {str(e)}")


@router.get("/api/mkdocs/stats/overview")
async def get_documentation_overview():
    """Get comprehensive documentation statistics overview."""
    try:
        published_dir = get_published_docs_dir()
        if not published_dir.exists():
            return {}

        stats = {
            "total_documents": 0,
            "by_type": defaultdict(int),
            "by_module": defaultdict(int),
            "by_owner": defaultdict(int),
            "by_status": defaultdict(int),
            "total_words": 0,
            "avg_words_per_doc": 0,
            "oldest_update": None,
            "newest_update": None,
            "total_atoms": set(),
        }

        atom_pattern = r"atom-[\w-]+"

        for md_file in published_dir.glob("*.md"):
            try:
                with open(md_file, "r", encoding="utf-8") as f:
                    content = f.read()

                frontmatter = parse_frontmatter(content)
                word_count = len(content.split())

                stats["total_documents"] += 1
                stats["total_words"] += word_count

                # Categorize by attributes
                stats["by_type"][frontmatter.get("template_type", "unknown")] += 1
                stats["by_module"][frontmatter.get("module", "unknown")] += 1
                stats["by_owner"][frontmatter.get("owner", "unknown")] += 1
                stats["by_status"][frontmatter.get("status", "unknown")] += 1

                # Track atoms
                atoms = set(re.findall(atom_pattern, content))
                stats["total_atoms"].update(atoms)

                # Track dates
                mod_time = datetime.fromtimestamp(md_file.stat().st_mtime)
                if stats["oldest_update"] is None or mod_time < stats["oldest_update"]:
                    stats["oldest_update"] = mod_time
                if stats["newest_update"] is None or mod_time > stats["newest_update"]:
                    stats["newest_update"] = mod_time

            except Exception as e:
                print(f"Error processing {md_file}: {e}")
                continue

        # Calculate averages
        if stats["total_documents"] > 0:
            stats["avg_words_per_doc"] = round(stats["total_words"] / stats["total_documents"], 1)

        # Convert sets and defaultdicts to regular dicts
        stats["total_atoms"] = len(stats["total_atoms"])
        stats["by_type"] = dict(stats["by_type"])
        stats["by_module"] = dict(stats["by_module"])
        stats["by_owner"] = dict(stats["by_owner"])
        stats["by_status"] = dict(stats["by_status"])

        # Format dates
        if stats["oldest_update"]:
            stats["oldest_update"] = stats["oldest_update"].isoformat()
        if stats["newest_update"]:
            stats["newest_update"] = stats["newest_update"].isoformat()

        return stats

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to get overview: {str(e)}")
